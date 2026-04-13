import os
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime


# ─────────────────────────────────────────────
# COLOUR / STYLE CONSTANTS
# ─────────────────────────────────────────────
PURPLE       = RGBColor(75, 0, 130)    # #4B0082
WHITE        = RGBColor(255, 255, 255)
GREEN        = RGBColor(0, 128, 0)
RED          = RGBColor(255, 0, 0)
PURPLE_HEX   = "4B0082"
LIGHT_PURPLE = "F3ECFA"               # screenshot block background


class BaseReport:
    """Shared ITSAR-styled DOCX formatting helpers for all clause reports."""

    def __init__(self, context, results):
        self.context = context
        self.results = results

    # ─────────────────────────────────────────
    # HEADING  (purple text + purple underline)
    # ─────────────────────────────────────────
    def add_itsar_heading(self, doc, text, level=1):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(2)

        run = para.add_run(text)
        run.bold            = True
        run.font.size       = Pt(16 if level == 1 else 14)
        run.font.color.rgb  = PURPLE

        pPr   = para._p.get_or_add_pPr()
        pBdr  = OxmlElement("w:pBdr")
        bot   = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "12")
        bot.set(qn("w:space"), "2")
        bot.set(qn("w:color"), PURPLE_HEX)
        pBdr.append(bot)
        pPr.append(pBdr)

        return para

    # ─────────────────────────────────────────
    # SUB-HEADING  (purple text, no underline)
    # ─────────────────────────────────────────
    def add_itsar_subheading(self, doc, text, level=2):
        para = doc.add_heading(text, level=level)
        run = para.runs[0]
        run.bold           = True
        run.font.size      = Pt(16 if level == 1 else 14)
        run.font.color.rgb = PURPLE
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(8)
        return para

    # ─────────────────────────────────────────
    # BOLD PARAGRAPH
    # ─────────────────────────────────────────
    def add_bold_paragraph(self, doc, text):
        p = doc.add_paragraph()
        p.add_run(text).bold = True
        return p

    # ─────────────────────────────────────────
    # KEEP-WITH-NEXT
    # ─────────────────────────────────────────
    def keep_with_next(self, para):
        para.paragraph_format.keep_with_next = True
        para.paragraph_format.keep_together  = True
        return para

    # ─────────────────────────────────────────
    # TABLE HEADER CELL  (purple bg, white bold)
    # ─────────────────────────────────────────
    def style_table_header(self, cell, color=PURPLE_HEX):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold           = True
                run.font.color.rgb = WHITE

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.top_margin    = Inches(0.15)
        cell.bottom_margin = Inches(0.15)
        cell.left_margin   = Inches(0.15)
        cell.right_margin  = Inches(0.15)

    # ─────────────────────────────────────────
    # PREVENT TABLE ROW SPLIT
    # ─────────────────────────────────────────
    def prevent_table_row_split(self, table):
        for row in table.rows:
            trPr      = row._tr.get_or_add_trPr()
            cantSplit = OxmlElement("w:cantSplit")
            trPr.append(cantSplit)

    # ─────────────────────────────────────────
    # SCREENSHOT EVIDENCE BLOCK (master.py style)
    # ─────────────────────────────────────────
    def add_screenshot_block(self, doc, title, image_path):
        TABLE_WIDTH = Inches(5.5)
        IMAGE_WIDTH = Inches(5.0)

        table = doc.add_table(rows=2, cols=1)
        table.alignment     = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False

        self.prevent_table_row_split(table)

        # Set explicit table width via XML to prevent overflow
        tblPr = table._tbl.tblPr
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(int(TABLE_WIDTH.emu / 635)))  # EMU to twips
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)

        table.columns[0].width = TABLE_WIDTH
        for row in table.rows:
            cell       = row.cells[0]
            cell.width = TABLE_WIDTH

            tcPr = cell._tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_PURPLE)
            tcPr.append(shd)

            # Constrain cell width explicitly
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(TABLE_WIDTH.emu / 635)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

            cell.top_margin    = Inches(0.15)
            cell.bottom_margin = Inches(0.15)
            cell.left_margin   = Inches(0.2)
            cell.right_margin  = Inches(0.2)

        # title cell
        title_cell = table.cell(0, 0)
        p_title    = title_cell.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.keep_with_next(p_title)

        run = p_title.add_run(title)
        run.bold           = True
        run.font.size      = Pt(11)
        run.font.color.rgb = PURPLE

        # image cell
        img_cell = table.cell(1, 0)
        p_img    = img_cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.keep_together = True
        p_img.add_run().add_picture(image_path, width=IMAGE_WIDTH)

        # purple border
        tblPr      = table._tbl.tblPr
        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"),   "single")
            border.set(qn("w:sz"),    "12")
            border.set(qn("w:color"), PURPLE_HEX)
            tblBorders.append(border)
        tblPr.append(tblBorders)

        return table

    # ─────────────────────────────────────────
    # GREY HORIZONTAL LINE
    # ─────────────────────────────────────────
    def add_grey_horizontal_line(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)

        p_elm = p._p
        p_pr  = p_elm.get_or_add_pPr()
        p_borders = OxmlElement("w:pBdr")

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "BFBFBF")

        p_borders.append(bottom)
        p_pr.append(p_borders)
        return p

    # ─────────────────────────────────────────
    # PAGE NUMBER FOOTER
    # ─────────────────────────────────────────
    def add_page_number(self, doc):
        footer    = doc.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.clear()

        for fld_type in ("begin", None, "end"):
            r = paragraph.add_run()
            if fld_type in ("begin", "end"):
                fc = OxmlElement("w:fldChar")
                fc.set(qn("w:fldCharType"), fld_type)
                r._r.append(fc)
            else:
                it = OxmlElement("w:instrText")
                it.text = "PAGE"
                r._r.append(it)

    # ─────────────────────────────────────────
    # TITLE
    # ─────────────────────────────────────────
    def add_title(self, doc, title_text="SECURITY TEST REPORT"):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = title.add_run(title_text)
        run.bold           = True
        run.font.size      = Pt(26)
        run.font.color.rgb = PURPLE

        doc.add_paragraph()

    # ─────────────────────────────────────────
    # DATA CELL PADDING
    # ─────────────────────────────────────────
    def add_data_cell_padding(self, table, skip_first_row=True):
        start = 1 if skip_first_row else 0
        for row in table.rows[start:]:
            for cell in row.cells:
                cell.top_margin    = Inches(0.12)
                cell.bottom_margin = Inches(0.12)
                cell.left_margin   = Inches(0.12)
                cell.right_margin  = Inches(0.12)

    # ─────────────────────────────────────────
    # DUT DETAILS TABLE
    # ─────────────────────────────────────────
    def add_dut_details(self, doc, context, section_num="1"):
        self.add_itsar_heading(doc, f"{section_num}. DUT Details", 2)

        rows_data = [
            ("Device",           getattr(context, "dut_model", "N/A") or "N/A"),
            ("Serial Number",    getattr(context, "dut_serial", "N/A") or "N/A"),
            ("Firmware Version", getattr(context, "dut_firmware", "N/A") or "N/A"),
            ("DUT IP Address",   getattr(context, "dut_ip", "N/A") or "N/A"),
        ]

        table = doc.add_table(rows=len(rows_data) + 1, cols=2)
        table.style = "Table Grid"

        for i, header in enumerate(["Parameter", "Value"]):
            cell      = table.rows[0].cells[i]
            cell.text = header
            self.style_table_header(cell)

        for i, (key, val) in enumerate(rows_data, start=1):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(val)

        self.add_data_cell_padding(table)
        self.prevent_table_row_split(table)

    # ─────────────────────────────────────────
    # RESULT SUMMARY TABLE
    # ─────────────────────────────────────────
    def add_result_summary(self, doc, results, section_num="9"):
        h = self.add_itsar_heading(doc, f"{section_num}. Test Case Result Summary", 2)
        self.keep_with_next(h)

        headers = ["SL No", "Test Case Name", "PASS/FAIL", "Remarks"]
        table   = doc.add_table(rows=len(results) + 1, cols=len(headers))
        table.style = "Table Grid"

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self.style_table_header(cell)

        for i, tc in enumerate(results, start=1):
            table.rows[i].cells[0].text = str(i)
            table.rows[i].cells[1].text = tc.name
            table.rows[i].cells[2].text = tc.status
            table.rows[i].cells[3].text = getattr(tc, "remarks", "")

            status_cell = table.rows[i].cells[2]
            for para in status_cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = GREEN if tc.status.upper() == "PASS" else RED

        self.add_data_cell_padding(table, skip_first_row=True)
        self.prevent_table_row_split(table)

    # ─────────────────────────────────────────
    # FIND SCREENSHOTS for a test case
    # ─────────────────────────────────────────
    def find_screenshots(self, clause, testcase_name):
        """Find all screenshots for a test case from the latest timestamp folder."""
        base = os.path.join(self.context.evidence.run_dir, str(clause), testcase_name)
        if not os.path.isdir(base):
            return []

        # Use the current run's timestamp
        ts_dir = os.path.join(base, self.context.evidence.date_prefix, "screenshots")
        if os.path.isdir(ts_dir):
            files = sorted(glob.glob(os.path.join(ts_dir, "*.png")))
            return files

        # Fallback: find the most recent timestamp folder
        timestamps = sorted(os.listdir(base), reverse=True)
        for ts in timestamps:
            ss_dir = os.path.join(base, ts, "screenshots")
            if os.path.isdir(ss_dir):
                files = sorted(glob.glob(os.path.join(ss_dir, "*.png")))
                if files:
                    return files
        return []

    # ─────────────────────────────────────────
    # REPORT OUTPUT PATH
    # ─────────────────────────────────────────
    def get_report_path(self, clause):
        """Returns output/{clause}/reports/{timestamp}_report.docx"""
        reports_dir = os.path.join(
            self.context.evidence.run_dir, str(clause), "reports"
        )
        os.makedirs(reports_dir, exist_ok=True)
        filename = f"{self.context.evidence.date_prefix}_report.docx"
        return os.path.join(reports_dir, filename)

    # ─────────────────────────────────────────
    # EMBED ALL SCREENSHOTS FOR A TEST CASE
    # ─────────────────────────────────────────
    def embed_testcase_screenshots(self, doc, clause, testcase_name, label_prefix=""):
        """Find and embed all screenshots for a test case, with explanations."""
        from reporting.pdf_base import describe_screenshot, _classify_port_for_desc

        screenshots = self.find_screenshots(clause, testcase_name)
        for img_path in screenshots:
            basename = os.path.splitext(os.path.basename(img_path))[0]
            # Clean up the timestamp prefix from filename for title
            parts = basename.split("_", 3)
            if len(parts) >= 4:
                title = f"{label_prefix}{parts[3]}" if label_prefix else parts[3]
            else:
                title = f"{label_prefix}{basename}" if label_prefix else basename

            # Enrich title for port-specific screenshots with service name
            lower = basename.lower()
            for proto in ("tcp", "udp", "sctp"):
                tag = f"{proto}_port_"
                if tag in lower:
                    port_str = lower.split(tag)[-1].split("_")[0]
                    if port_str.isdigit():
                        svc, url, common = _classify_port_for_desc(int(port_str))
                        verdict_tag = "PASS" if common else "FAIL"
                        title = (f"{label_prefix}{proto.upper()} Port {port_str} "
                                 f"— {svc} [{verdict_tag}]")
                    break

            self.add_screenshot_block(doc, title, img_path)

            # Add Observations paragraph below the image
            desc = describe_screenshot(img_path)
            if desc:
                obs_heading = doc.add_paragraph()
                obs_run = obs_heading.add_run("Observations:")
                obs_run.bold = True
                obs_run.font.size = Pt(10)
                obs_run.font.color.rgb = PURPLE

                obs_para = doc.add_paragraph()
                obs_text = obs_para.add_run(desc)
                obs_text.italic = True
                obs_text.font.size = Pt(9)
                obs_text.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)  # grey
            doc.add_paragraph()

    # ─────────────────────────────────────────
    # FRONT PAGE  (title, metadata, overall result)
    # ─────────────────────────────────────────
    def _add_front_page(self, doc, context, results):
        """Professional front page with title, metadata table, and overall result banner."""

        # Compute overall result
        failed = [tc for tc in results if getattr(tc, "status", "FAIL").upper() != "PASS"]
        overall = "PASS" if not failed else "FAIL"

        # ── Title ──
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("SECURITY TEST REPORT")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = PURPLE

        # ── Subtitle ──
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = sub.add_run(f"ITSAR Compliance Test Report -- Clause {context.clause}")
        run2.bold = True
        run2.font.size = Pt(14)
        run2.font.color.rgb = PURPLE
        doc.add_paragraph()

        # ── Metadata table ──
        now = datetime.now()
        meta_rows = [
            ("ITSAR Clause",     str(context.clause)),
            ("DUT Name",         getattr(context, "dut_model", "N/A") or "N/A"),
            ("DUT IP Address",   getattr(context, "dut_ip", "N/A") or "N/A"),
            ("DUT Firmware",     getattr(context, "dut_firmware", "N/A") or "N/A"),
            ("Test Date",        now.strftime("%Y-%m-%d")),
            ("Test Time",        now.strftime("%H:%M:%S")),
            ("Total Test Cases", str(len(results))),
            ("Overall Result",   overall),
        ]

        table = doc.add_table(rows=len(meta_rows) + 1, cols=2)
        table.style = "Table Grid"

        for i, header in enumerate(["Parameter", "Value"]):
            cell = table.rows[0].cells[i]
            cell.text = header
            self.style_table_header(cell)

        for i, (key, val) in enumerate(meta_rows, start=1):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(val)

            # Color the overall result cell
            if key == "Overall Result":
                cell = table.rows[i].cells[1]
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = GREEN if overall == "PASS" else RED

        self.add_data_cell_padding(table)
        self.prevent_table_row_split(table)

        doc.add_paragraph()

        # ── Overall result banner ──
        banner = doc.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_b = banner.add_run(f"OVERALL RESULT:  {overall}")
        run_b.bold = True
        run_b.font.size = Pt(18)
        run_b.font.color.rgb = GREEN if overall == "PASS" else RED

        doc.add_page_break()

    # ─────────────────────────────────────────
    # COMPLIANCE ANALYSIS TABLE
    # ─────────────────────────────────────────
    def _add_compliance_analysis(self, doc, results, section_num="11"):
        """Add compliance analysis table mapping requirements to test results."""
        h = self.add_itsar_heading(doc, f"{section_num}. Compliance Analysis", 2)
        self.keep_with_next(h)

        headers = ["Clause Requirement", "Test Case(s)", "Result"]
        table = doc.add_table(rows=len(results) + 1, cols=len(headers))
        table.style = "Table Grid"

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self.style_table_header(cell)

        for i, tc in enumerate(results, start=1):
            desc = getattr(tc, "description", getattr(tc, "name", f"TC{i}"))
            table.rows[i].cells[0].text = str(desc)[:70]
            table.rows[i].cells[1].text = getattr(tc, "name", f"TC{i}")
            table.rows[i].cells[2].text = getattr(tc, "status", "N/A")

            status_cell = table.rows[i].cells[2]
            for para in status_cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    status = getattr(tc, "status", "N/A")
                    run.font.color.rgb = GREEN if status.upper() == "PASS" else RED

        self.add_data_cell_padding(table, skip_first_row=True)
        self.prevent_table_row_split(table)

    # ─────────────────────────────────────────
    # CONCLUSION & RECOMMENDATIONS
    # ─────────────────────────────────────────
    def _add_conclusion(self, doc, results, section_num="12", recommendations=None):
        """Add conclusion section with optional recommendations."""
        self.add_itsar_heading(doc, f"{section_num}. Conclusion", 2)

        total = len(results)
        failed = [tc for tc in results if getattr(tc, "status", "FAIL").upper() != "PASS"]

        if not failed:
            doc.add_paragraph(
                f"All {total} test case(s) passed. The DUT complies with the "
                "requirements specified in this ITSAR clause."
            )
        else:
            names = ", ".join(getattr(tc, "name", "?") for tc in failed)
            doc.add_paragraph(
                f"The test run identified {len(failed)} failing test case(s): {names}. "
                "The DUT does NOT fully comply with this ITSAR clause. "
                "Remediation is required before the DUT can be considered compliant."
            )

        if recommendations:
            doc.add_paragraph()
            h = self.add_itsar_subheading(doc, f"{section_num}.1 Recommendations", 2)
            self.keep_with_next(h)
            for rec in recommendations:
                doc.add_paragraph(f"\u2022 {rec}")
