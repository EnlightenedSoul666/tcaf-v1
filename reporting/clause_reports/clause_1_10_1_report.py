import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reporting.base_report import BaseReport, GREEN, RED, PURPLE, PURPLE_HEX, WHITE


# ===========================================================================
#  Screenshot grouping helpers for sub-testcase sections
# ===========================================================================

def _extract_icmp_type(filename):
    """Extract ICMP type number from a screenshot filename."""
    match = re.search(r'type_(\d+)', filename.lower())
    return int(match.group(1)) if match else None


def _routing_sort_key(path):
    """Sort routing screenshots: before -> table -> after."""
    name = os.path.basename(path).lower()
    if "before" in name:
        return 0
    if "routing_table" in name:
        return 1
    if "after" in name:
        return 2
    return 3


def _redirect_sort_key(path):
    """Sort redirect screenshots: before -> ping -> after -> pcap -> wireshark."""
    name = os.path.basename(path).lower()
    if "redirect_before" in name:
        return 0
    if "redirect_ping" in name:
        return 1
    if "redirect_after" in name:
        return 2
    if "redirect_pcap" in name:
        return 3
    # Wireshark packet screenshot
    if "packet_frame" in name or "redirect_packet" in name:
        return 4
    return 5


def _process_sort_key(path):
    """Sort RS/RA process screenshots: before -> after."""
    name = os.path.basename(path).lower()
    if "process_before" in name:
        return 0
    if "process_after" in name:
        return 1
    return 2


def _screenshot_label(path):
    """Generate a concise human-readable label for a screenshot."""
    name = os.path.basename(path).lower()

    # Routing
    if "traceroute" in name and "before" in name:
        return "Traceroute BEFORE route setup"
    if "traceroute" in name and "after" in name:
        return "Traceroute AFTER route setup"
    if "routing_table" in name:
        return "Routing table (ip route show)"

    # Redirect sub-steps (chronological: setup -> stimulus -> evidence)
    if "redirect_before" in name:
        return "Step 1 — Traceroute BEFORE redirect (path forced through DuT)"
    if "redirect_ping" in name:
        return "Step 2 — Ping auxiliary to provoke ICMP Redirect from DuT"
    if "redirect_after" in name:
        return "Step 3 — Traceroute AFTER redirect (compliance evidence)"
    if "redirect_pcap" in name:
        return "Step 4 — tshark shows the Redirect packet in the PCAP"
    if "redirect_packet" in name or ("packet_frame" in name and "redirect" in name):
        return "Step 5 — Wireshark frame detail of the captured Redirect"

    # Process sub-steps (RS/RA traceroute comparison)
    if "process_before" in name:
        if "type_133" in name:
            return "Traceroute BEFORE sending Router Solicitation (baseline path)"
        if "type_134" in name:
            return "Traceroute BEFORE sending Router Advertisement (baseline path)"
        return "Traceroute BEFORE sending crafted packet (baseline path)"
    if "process_after" in name:
        if "type_133" in name:
            return "Traceroute AFTER Router Solicitation (should match baseline)"
        if "type_134" in name:
            return "Traceroute AFTER Router Advertisement (should match baseline)"
        return "Traceroute AFTER sending crafted packet (should match baseline)"

    # Send / Not Permitted — terminal tshark
    if name.startswith("tester") or (not name.startswith("packet_frame")):
        if "notpermitted" in name:
            return "Terminal: tshark check for NOT PERMITTED response"
        if "send" in name:
            return "Terminal: tshark analysis of DuT response"

    # Wireshark packet screenshots
    if "packet_frame" in name:
        return "Wireshark: packet capture detail"

    return os.path.splitext(os.path.basename(path))[0]


def _group_screenshots(screenshots, sub_results, context):
    """
    Group screenshots into logical sub-test sections.
    Returns [(title, [screenshot_paths], observation, status), ...].
    """
    routing = []
    by_send_type = {}       # icmp_type -> [paths]
    by_redirect = []
    by_process_type = {}    # icmp_type -> [paths]
    uncategorized = []

    for ss in screenshots:
        name = os.path.basename(ss).lower()
        icmp_type = _extract_icmp_type(name)

        if "traceroute" in name and "redirect" not in name and "process" not in name:
            routing.append(ss)
        elif "routing_table" in name:
            routing.append(ss)
        elif "redirect" in name:
            by_redirect.append(ss)
        elif "process" in name:
            if icmp_type is not None:
                by_process_type.setdefault(icmp_type, []).append(ss)
            else:
                uncategorized.append(ss)
        elif "notpermitted" in name:
            if icmp_type is not None:
                by_send_type.setdefault(icmp_type, []).append(ss)
            else:
                uncategorized.append(ss)
        elif "send" in name:
            if icmp_type is not None:
                by_send_type.setdefault(icmp_type, []).append(ss)
            else:
                uncategorized.append(ss)
        elif "packet_frame" in name:
            # Wireshark screenshots — match to the right group
            if "redirect" in name:
                by_redirect.append(ss)
            elif "process" in name and icmp_type is not None:
                by_process_type.setdefault(icmp_type, []).append(ss)
            elif icmp_type is not None:
                by_send_type.setdefault(icmp_type, []).append(ss)
            else:
                uncategorized.append(ss)
        else:
            uncategorized.append(ss)

    groups = []

    # 1. Routing setup
    if routing:
        routing.sort(key=_routing_sort_key)
        v = "IPv6" if any("ipv6" in os.path.basename(s).lower() for s in routing) else "IPv4"
        openwrt = context.openwrt_ip or "OpenWRT"
        nonsense = context.nonsense_ip or context.nonsense_ipv6 or "nonsense IP"
        aux = getattr(context, "auxiliary_ip", None) or getattr(context, "auxiliary_ipv6", None) or "auxiliary machine"
        if v == "IPv6":
            openwrt = context.openwrt_ipv6 or openwrt
            nonsense = context.nonsense_ipv6 or nonsense
            aux = getattr(context, "auxiliary_ipv6", None) or aux
        obs = (
            f"Static routes are added so packets to {nonsense} and {aux} "
            f"are forwarded through the DuT (OpenWRT at {openwrt}) instead "
            f"of taking the default direct path. The BEFORE traceroute shows "
            f"the default network path. The routing table confirms the new "
            f"routes are in place. The AFTER traceroute verifies packets now "
            f"traverse through the router as an intermediate hop. This is "
            f"required to trigger ICMP errors (Destination Unreachable, Time "
            f"Exceeded, Redirect) from the DuT."
        )
        groups.append((f"Routing Setup ({v})", routing, obs, None))

    # 2. Send tests (in sub_results order)
    for sr in sub_results:
        if sr.get("test_type") != "Send":
            continue
        icmp_type = sr["icmp_type"]
        ss_list = sorted(by_send_type.get(icmp_type, []))
        if not ss_list:
            continue
        category = sr.get("category", "Permitted")
        if category == "Not Permitted":
            title = f"{sr['icmp_name']} — NOT PERMITTED"
        else:
            title = f"{sr['icmp_name']} — {category}"
        groups.append((title, ss_list, sr.get("description", ""), sr.get("status")))

    # 3. Process tests (in sub_results order)
    for sr in sub_results:
        if sr.get("test_type") != "Process":
            continue
        icmp_type = sr["icmp_type"]
        if icmp_type in (5, 137):
            ss_list = sorted(by_redirect, key=_redirect_sort_key)
        else:
            ss_list = sorted(by_process_type.get(icmp_type, []),
                             key=_process_sort_key)
        if not ss_list:
            continue
        title = f"Process Test: {sr['icmp_name']} (Type {icmp_type}) — NOT PERMITTED"
        groups.append((title, ss_list, sr.get("description", ""), sr.get("status")))

    # 4. Uncategorized
    if uncategorized:
        groups.append(("Additional Evidence", sorted(uncategorized), "", None))

    return groups


# ===========================================================================
#  DOCX Report
# ===========================================================================

class Clause1101Report(BaseReport):
    """ITSAR Clause 1.10.1 -- ICMP Type Filtering Compliance Report (DOCX)."""

    CLAUSE_ID = "1.10.1"

    def _add_icmp_reference_tables(self, doc):
        """Add the ETSI ICMP type reference tables."""

        # -- Permitted Types --
        self.add_itsar_subheading(doc, "Permitted ICMP Types (ETSI TS 133 117)", 2)
        permitted = [
            ("IPv4", "IPv6", "Name",                    "Send",      "Respond To"),
            ("0",    "128",  "Echo Reply",               "Optional",  "N/A"),
            ("3",    "1",    "Destination Unreachable",   "Permitted", "N/A"),
            ("8",    "129",  "Echo Request",              "Permitted", "Optional"),
            ("11",   "3",    "Time Exceeded",             "Optional",  "N/A"),
            ("12",   "4",    "Parameter Problem",         "Permitted", "N/A"),
            ("N/A",  "2",    "Packet Too Big",            "Permitted", "N/A"),
            ("N/A",  "135",  "Neighbour Solicitation",    "Permitted", "Permitted"),
            ("N/A",  "136",  "Neighbour Advertisement",   "Permitted", "N/A"),
        ]
        table = doc.add_table(rows=len(permitted), cols=5)
        table.style = "Table Grid"
        for i, header in enumerate(permitted[0]):
            cell = table.rows[0].cells[i]
            cell.text = header
            self.style_table_header(cell)
        for r, row_data in enumerate(permitted[1:], start=1):
            for c, val in enumerate(row_data):
                table.rows[r].cells[c].text = val
        self.add_data_cell_padding(table, skip_first_row=True)
        self.prevent_table_row_split(table)
        doc.add_paragraph()

        # -- Not Permitted Types --
        self.add_itsar_subheading(doc, "Not Permitted ICMP Types", 2)
        not_permitted = [
            ("IPv4", "IPv6", "Name",                   "Send",          "Respond To",      "Process"),
            ("5",    "137",  "Redirect",                "N/A",           "N/A",             "Not Permitted"),
            ("13",   "N/A",  "Timestamp Request",       "N/A",           "Not Permitted",   "N/A"),
            ("14",   "N/A",  "Timestamp Reply",         "Not Permitted", "N/A",             "N/A"),
            ("N/A",  "133",  "Router Solicitation",     "N/A",           "Not Permitted",   "Not Permitted"),
            ("N/A",  "134",  "Router Advertisement",    "N/A",           "N/A",             "Not Permitted"),
        ]
        table2 = doc.add_table(rows=len(not_permitted), cols=6)
        table2.style = "Table Grid"
        for i, header in enumerate(not_permitted[0]):
            cell = table2.rows[0].cells[i]
            cell.text = header
            self.style_table_header(cell)
        for r, row_data in enumerate(not_permitted[1:], start=1):
            for c, val in enumerate(row_data):
                table2.rows[r].cells[c].text = val
        self.add_data_cell_padding(table2, skip_first_row=True)
        self.prevent_table_row_split(table2)
        doc.add_paragraph()

    def _add_observation(self, doc, text):
        """Add an Observations block with purple heading and grey italic text."""
        obs_heading = doc.add_paragraph()
        obs_run = obs_heading.add_run("Observations:")
        obs_run.bold = True
        obs_run.font.size = Pt(10)
        obs_run.font.color.rgb = PURPLE

        obs_para = doc.add_paragraph()
        obs_text = obs_para.add_run(text)
        obs_text.italic = True
        obs_text.font.size = Pt(9)
        obs_text.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

    def _add_status_badge(self, doc, status):
        """Add a status line (PASS/FAIL) with appropriate color."""
        p = doc.add_paragraph("Status: ")
        run = p.add_run(status)
        run.bold = True
        run.font.color.rgb = GREEN if status.upper() == "PASS" else RED

    def generate(self, context, results):
        doc = Document()

        self.add_page_number(doc)
        self.add_title(doc)

        # -- Front Page --
        self._add_front_page(doc, context, results)

        # -- 1. DUT Details --
        self.add_dut_details(doc, context, section_num="1")
        doc.add_paragraph()

        # -- 2. ITSAR Information --
        self.add_itsar_heading(doc, "2. ITSAR Information", 2)
        table = doc.add_table(rows=5, cols=2)
        table.style = "Table Grid"
        for i, header in enumerate(["Field", "Value"]):
            cell = table.rows[0].cells[i]
            cell.text = header
            self.style_table_header(cell)
        info_rows = [
            ("ITSAR Section",   "1.10.1"),
            ("Requirement",     "ICMP Type Filtering"),
            ("ETSI Reference",  "ETSI TS 133 117 V17.2.0, Section 4.2.4.1.1.2"),
            ("TSDSI Reference", "TSDSI STD T1.3GPP 33.117-17.2.0 V.1.0.0"),
        ]
        for r, (k, v) in enumerate(info_rows, start=1):
            table.rows[r].cells[0].text = k
            table.rows[r].cells[1].text = v
        self.add_data_cell_padding(table)
        self.prevent_table_row_split(table)
        doc.add_paragraph()

        # -- 3. Requirement Description --
        self.add_itsar_heading(doc, "3. Requirement Description", 2)
        doc.add_paragraph(
            "Processing of ICMPv4 and ICMPv6 packets which are not required for operation "
            "shall be disabled on the CPE. In particular, there are certain types of ICMPv4 and "
            "ICMPv6 that are not used in most networks, but represent a risk. Refer standards "
            "such as RFC 6192, RFC 7279, RFC 4890."
        )
        doc.add_paragraph()

        # ICMP Reference Tables
        self._add_icmp_reference_tables(doc)

        # -- 4. Preconditions --
        self.add_itsar_heading(doc, "4. Preconditions", 2)
        for item in [
            "The tester system has network connectivity to the DUT.",
            "The DUT IPv4 address is reachable from the tester.",
            "The DUT IPv6 address is reachable from the tester (if applicable).",
            "The tester has Scapy, tcpdump, tshark, and Wireshark installed.",
            "The tester has root/sudo privileges for raw packet injection.",
            "For Process tests: a third host and routing configuration are available.",
        ]:
            doc.add_paragraph(f"\u2022 {item}")
        doc.add_paragraph()

        # -- 5. Test Objective --
        self.add_itsar_heading(doc, "5. Test Objective", 2)
        doc.add_paragraph(
            "To verify that the DUT correctly filters ICMP and ICMPv6 packet types "
            "according to the ITSAR requirements. The tester sends various ICMP type "
            "packets to the DUT and captures both the request and the response to "
            "determine whether the DUT allows or blocks each type."
        )
        doc.add_paragraph()

        # -- 6. Test Scenario --
        self.add_itsar_heading(doc, "6. Test Scenario", 2)

        self.add_itsar_subheading(doc, "6.1 Network Fundamentals", 2)

        self.add_bold_paragraph(doc, "IPv4 and IPv6 Subnet Basics")
        doc.add_paragraph(
            "All devices in this test topology reside on the same IPv4 private subnet "
            "(10.208.207.0/24) and share an IPv6 Unique Local Address (ULA) prefix "
            "(fdd4:48ab:15e6::/60). IPv4 addresses are manually configured (static) or "
            "assigned via DHCP from the OpenWRT router. IPv6 addresses are auto-configured "
            "via SLAAC (Stateless Address Auto-Configuration): the router sends Router "
            "Advertisements containing the subnet prefix, and each host derives its own "
            "address by combining the prefix with its MAC-based interface identifier."
        )

        self.add_bold_paragraph(doc, "Router Solicitation and Router Advertisement")
        doc.add_paragraph(
            "When an IPv6-enabled host boots, it sends a Router Solicitation (ICMPv6 Type 133) "
            "to discover on-link routers; the router replies with a Router Advertisement "
            "(ICMPv6 Type 134) containing prefix, MTU, and default-gateway information. "
            "This exchange enables zero-configuration IPv6 networking, unlike IPv4 which "
            "requires either a DHCP server or manual static configuration."
        )

        self.add_bold_paragraph(doc, "Routing Setup for ICMP Tests")
        doc.add_paragraph(
            "To test Send and Process categories, the tester adds static routes so that "
            "packets destined for the auxiliary machine and the nonsense IP travel through "
            "the OpenWRT router instead of directly to the target. This forces OpenWRT to "
            "generate ICMP errors (Destination Unreachable, Time Exceeded, Redirect) that "
            "would not occur on a flat L2 segment. Traceroutes are captured before and after "
            "route setup to evidence the path change."
        )

        self.add_bold_paragraph(doc, "How ICMP Redirect Is Triggered")
        doc.add_paragraph(
            "An ICMP Redirect (Type 5 / ICMPv6 Type 137) is generated by a router when it "
            "receives a packet and the best next-hop for that destination is on the same "
            "interface the packet arrived on. The router forwards the packet but sends a "
            "Redirect back to the sender, advising it to send future packets directly to "
            "the better gateway. If the DUT processes this Redirect, its routing table changes "
            "silently -- a security risk that ETSI prohibits."
        )

        self.add_bold_paragraph(doc, "Caution with ip route Commands")
        doc.add_paragraph(
            "Incorrect static routes can black-hole traffic or create routing loops. "
            "For example, adding a route with a gateway that is itself unreachable will cause "
            "all matching packets to be silently dropped. Adding overlapping routes without "
            "proper metric values can lead to unpredictable path selection. The test framework "
            "always cleans up (deletes) added routes in the teardown phase to prevent residual "
            "misrouting."
        )

        self.add_bold_paragraph(doc, "VM Boot Order")
        doc.add_paragraph(
            "The host VM (Kali tester) should be started first, followed by OpenWRT. "
            "This ensures Kali obtains its IP address from the host machine's Windows network "
            "(via the bridged adapter) rather than from OpenWRT's DHCP pool. If OpenWRT boots "
            "first, Kali may receive its default gateway from OpenWRT instead of from the "
            "Windows host, which can isolate the tester from the external network."
        )
        doc.add_paragraph()

        self.add_itsar_subheading(doc, "6.2 Tools Required", 2)
        for tool in ["Scapy (ICMP packet crafting)", "tcpdump (packet capture)",
                     "tshark (packet analysis)", "Wireshark (visual evidence)",
                     "Linux-based tester system"]:
            doc.add_paragraph(f"\u2022 {tool}")

        self.add_itsar_subheading(doc, "6.3 Test Execution Steps", 2)
        for step in [
            "Set up routing so packets to nonsense/auxiliary IPs go through the DuT (OpenWRT).",
            "Start packet capture on the tester interface using tcpdump.",
            "Send each purposeful ICMP packet (ping, TTL=1, malformed options, etc.) to trigger a specific response.",
            "Wait for DuT responses and stop the packet capture.",
            "Analyze the captured PCAP using tshark to verify each expected response.",
            "Take terminal and Wireshark screenshots with educational descriptions for each test.",
            "For Not Permitted tests: verify the DuT does NOT send the forbidden response type.",
            "For Process tests: send Redirect/RS/RA and verify no routing configuration change.",
        ]:
            doc.add_paragraph(f"\u2022 {step}")
        doc.add_paragraph()

        # -- 7. Expected Results --
        self.add_itsar_heading(doc, "7. Expected Results", 2)
        doc.add_paragraph(
            "The DUT should respond to allowed ICMP types (e.g., Echo Request/Reply) "
            "and block or ignore disallowed ICMP types. The captured packets should "
            "show the expected filtering behavior for each ICMP type tested."
        )
        doc.add_paragraph()

        # ==================================================================
        # 8. Test Execution — with sub-testcase sections
        # ==================================================================
        self.add_itsar_heading(doc, "8. Test Execution", 2)

        for idx, tc in enumerate(results, start=1):
            # -- 8.{idx} Test Case header --
            h = self.add_itsar_subheading(doc, f"8.{idx} Test Case: {tc.name}", 2)
            self.keep_with_next(h)
            doc.add_paragraph(f"Description: {tc.description}")

            p = doc.add_paragraph("Overall Result: ")
            run = p.add_run(tc.status)
            run.bold = True
            run.font.color.rgb = GREEN if tc.status.upper() == "PASS" else RED

            self.add_grey_horizontal_line(doc)

            # Group screenshots by sub-test
            screenshots = self.find_screenshots(self.CLAUSE_ID, tc.name)
            sub_results = getattr(tc, "sub_results", [])
            groups = _group_screenshots(screenshots, sub_results, context)

            if not groups:
                # Fallback: embed screenshots ungrouped (old behavior)
                self.embed_testcase_screenshots(
                    doc, self.CLAUSE_ID, tc.name, label_prefix=f"TC{idx} -- ")
                continue

            for sub_idx, (title, group_ss, observation, status) in enumerate(groups, start=1):
                # -- 8.{idx}.{sub_idx} Sub-test heading --
                self.add_itsar_subheading(
                    doc, f"8.{idx}.{sub_idx} {title}", 3)

                if status:
                    self._add_status_badge(doc, status)

                # Embed each screenshot with a concise label
                for img_path in group_ss:
                    label = _screenshot_label(img_path)
                    self.add_screenshot_block(doc, label, img_path)

                # Observation below the screenshots
                if observation:
                    self._add_observation(doc, observation)

                doc.add_paragraph()

        doc.add_paragraph()

        # -- 9. Test Observation --
        self.add_itsar_heading(doc, "9. Test Observation", 2)
        failed = [tc for tc in results if tc.status.upper() != "PASS"]
        if failed:
            names = ", ".join(tc.name for tc in failed)
            doc.add_paragraph(
                f"The following test case(s) did not pass: {names}. "
                "This indicates that the DUT may not be filtering all ICMP types "
                "as required by the ITSAR specification."
            )
        else:
            doc.add_paragraph(
                "All ICMP type filtering test cases passed successfully. "
                "The DUT correctly handles the tested ICMP and ICMPv6 packet types "
                "in accordance with the ITSAR requirements."
            )
        doc.add_paragraph()

        # -- 10. Test Case Result Summary --
        self.add_result_summary(doc, results, section_num="10")
        doc.add_paragraph()

        # Per-ICMP-type detailed results table
        all_sub = []
        for tc in results:
            for sr in getattr(tc, "sub_results", []):
                all_sub.append(sr)

        if all_sub:
            self.add_itsar_subheading(doc, "10.1 Per-ICMP-Type Detailed Results", 2)
            doc.add_paragraph(
                "The table below shows the individual PASS/FAIL result for each ICMP type "
                "tested across all categories (Send, Process)."
            )

            all_sub.sort(key=lambda s: (s.get("ip_version", 0),
                                         s.get("test_type", ""),
                                         s.get("icmp_type", 0)))

            headers = ["ICMP Type", "Name", "IP Ver", "Category", "Restriction", "PASS/FAIL"]
            table = doc.add_table(rows=len(all_sub) + 1, cols=len(headers))
            table.style = "Table Grid"

            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                self.style_table_header(cell)

            for i, sr in enumerate(all_sub, start=1):
                table.rows[i].cells[0].text = str(sr.get("icmp_type", "?"))
                table.rows[i].cells[1].text = sr.get("icmp_name", "Unknown")
                table.rows[i].cells[2].text = f"IPv{sr.get('ip_version', '?')}"
                table.rows[i].cells[3].text = sr.get("test_type", "?")
                table.rows[i].cells[4].text = sr.get("category", "?")
                table.rows[i].cells[5].text = sr.get("status", "N/A")

                # Color the status cell
                status_cell = table.rows[i].cells[5]
                for para in status_cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = GREEN if sr.get("status") == "PASS" else RED

            self.add_data_cell_padding(table, skip_first_row=True)
            self.prevent_table_row_split(table)
            doc.add_paragraph()

        # -- 11. Compliance Analysis --
        self._add_compliance_analysis(doc, results, section_num="11")
        doc.add_paragraph()

        # -- 12. Conclusion & Recommendations --
        self._add_conclusion(doc, results, section_num="12", recommendations=[
            "Configure the DUT firewall to drop ICMPv4 Type 5 (Redirect), 13 (Timestamp), and 14 (Timestamp Reply).",
            "Configure the DUT to drop ICMPv6 Type 133 (Router Solicitation), 134 (Router Advertisement), and 137 (Redirect).",
            "Ensure the DUT does not originate Timestamp Reply (Type 14) packets.",
            "Verify that processing of Redirect, RS, and RA packets is disabled.",
            "Re-run this test suite after any firewall or routing configuration change.",
        ])

        # Save
        report_path = self.get_report_path(self.CLAUSE_ID)
        doc.save(report_path)
        return report_path
