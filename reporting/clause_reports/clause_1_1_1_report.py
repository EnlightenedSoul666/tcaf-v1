import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reporting.base_report import BaseReport, GREEN, RED, PURPLE


class Clause111Report(BaseReport):
    """ITSAR Clause 1.1.1 -- CPE Authentication Compliance Report (DOCX)."""

    CLAUSE_ID = "1.1.1"

    def generate(self, context, results):
        doc = Document()

        self.add_page_number(doc)
        self.add_title(doc)

        # ── Front Page ──
        self._add_front_page(doc, context, results)

        # ── 1. DUT Details ──
        self.add_dut_details(doc, context, section_num="1")
        doc.add_paragraph()

        # ── 2. ITSAR Information ──
        self.add_itsar_heading(doc, "2. ITSAR Information", 2)
        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        for i, header in enumerate(["Field", "Value"]):
            cell = table.rows[0].cells[i]
            cell.text = header
            self.style_table_header(cell)
        table.rows[1].cells[0].text = "ITSAR Section"
        table.rows[1].cells[1].text = getattr(context, "itsar_section", "1.1.1")
        table.rows[2].cells[0].text = "Requirement"
        table.rows[2].cells[1].text = getattr(context, "itsar_requirement", "CPE Authentication")
        self.add_data_cell_padding(table)
        self.prevent_table_row_split(table)
        doc.add_paragraph()

        # ── 3. Requirement Description ──
        self.add_itsar_heading(doc, "3. Requirement Description", 2)
        doc.add_paragraph(
            "The CPE shall communicate with authenticated management entities only. "
            "The protocols used for the CPE management shall support mutual authentication "
            "mechanisms, preferably with pre-shared key arrangements or by equivalent entity "
            "mutual authentication mechanisms. This shall be verified for all protocols used "
            "for CPE management. (This feature shall be supported on all WAN management interfaces)."
        )
        doc.add_paragraph()

        # ── 4. Preconditions ──
        self.add_itsar_heading(doc, "4. Preconditions", 2)
        for item in [
            "The tester system has network connectivity to the DUT.",
            "SSH service is running on the DUT and accessible from the tester.",
            "The tester has nmap, tshark, Wireshark, and OpenSSH installed.",
            "HTTPS service is running on the DUT (if applicable).",
        ]:
            doc.add_paragraph(f"\u2022 {item}")
        doc.add_paragraph()

        # ── 5. Test Objective ──
        self.add_itsar_heading(doc, "5. Test Objective", 2)
        doc.add_paragraph(
            "To verify that the DUT management traffic is protected using secure "
            "cryptographic controls. The tester enumerates supported SSH/TLS algorithms, "
            "captures handshake traffic, verifies cipher strength against NIST/ITSAR "
            "requirements, and attempts weak algorithm negotiation to confirm rejection."
        )
        doc.add_paragraph()

        # ── 6. Test Scenario ──
        self.add_itsar_heading(doc, "6. Test Scenario", 2)
        self.add_itsar_subheading(doc, "6.1 Tools Required", 2)
        for tool in ["Nmap (SSH/TLS cipher enumeration)", "Wireshark / tshark",
                     "OpenSSH", "openssl s_client", "Linux-based tester system"]:
            doc.add_paragraph(f"\u2022 {tool}")

        self.add_itsar_subheading(doc, "6.2 Test Execution Steps", 2)
        for step in [
            "Enumerate SSH algorithms using nmap --script ssh2-enum-algos.",
            "Capture SSH handshake traffic and verify cipher/KEX/MAC algorithms.",
            "Attempt SSH connection with deliberately weak ciphers to verify rejection.",
            "If HTTPS is available, enumerate TLS ciphers and capture Server Hello.",
        ]:
            doc.add_paragraph(f"\u2022 {step}")
        doc.add_paragraph()

        # ── 7. Expected Results ──
        self.add_itsar_heading(doc, "7. Expected Results", 2)
        doc.add_paragraph(
            "The DUT shall support only NIST-recommended cryptographic algorithms. "
            "Attempts to connect with weak algorithms shall be rejected by the DUT."
        )
        doc.add_paragraph()

        # ── 8. Test Execution ──
        self.add_itsar_heading(doc, "8. Test Execution", 2)

        for idx, tc in enumerate(results, start=1):
            h = self.add_itsar_subheading(doc, f"8.{idx} Test Case: {tc.name}", 2)
            self.keep_with_next(h)

            doc.add_paragraph(f"Description: {tc.description}")

            p = doc.add_paragraph("Result: ")
            run = p.add_run(tc.status)
            run.bold = True
            run.font.color.rgb = GREEN if tc.status.upper() == "PASS" else RED

            self.add_grey_horizontal_line(doc)

            # Embed evidence from test case
            for evidence in tc.evidence:
                screenshot = evidence.get("screenshot") if isinstance(evidence, dict) else getattr(evidence, "screenshot", None)
                if screenshot and os.path.exists(screenshot):
                    self.add_screenshot_block(doc, f"Evidence: {os.path.basename(screenshot)}", screenshot)
                    doc.add_paragraph()

            self.embed_testcase_screenshots(doc, self.CLAUSE_ID, tc.name, label_prefix=f"TC{idx} -- ")

        doc.add_paragraph()

        # ── 9. Test Observation ──
        self.add_itsar_heading(doc, "9. Test Observation", 2)
        failed = [tc for tc in results if tc.status.upper() != "PASS"]
        if failed:
            names = ", ".join(tc.name for tc in failed)
            doc.add_paragraph(
                f"The following test case(s) did not pass: {names}. "
                "This indicates the DUT may permit insecure cryptographic configurations."
            )
        else:
            doc.add_paragraph(
                "All test cases passed. The DUT complies with the prescribed secure "
                "cryptographic requirements for management communication."
            )
        doc.add_paragraph()

        # ── 10. Test Case Result Summary ──
        self.add_result_summary(doc, results, section_num="10")
        doc.add_paragraph()

        # ── 11. Compliance Analysis ──
        self._add_compliance_analysis(doc, results, section_num="11")
        doc.add_paragraph()

        # ── 12. Conclusion & Recommendations ──
        self._add_conclusion(doc, results, section_num="12", recommendations=[
            "Ensure the DUT only enables NIST-recommended SSH ciphers, MAC, and KEX algorithms.",
            "Disable all weak/deprecated algorithms (DES, RC4, MD5, SHA1-based MACs, etc.).",
            "If HTTPS is available, enforce TLSv1.2+ with strong cipher suites only.",
            "Re-run this test suite after any configuration change to verify compliance.",
        ])

        # Save
        report_path = self.get_report_path(self.CLAUSE_ID)
        doc.save(report_path)
        return report_path
