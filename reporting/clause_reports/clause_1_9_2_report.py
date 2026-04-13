import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reporting.base_report import BaseReport, GREEN, RED, PURPLE

# IANA master registry — cited in the Port Classification table
IANA_URL = (
    "https://www.iana.org/assignments/service-names-port-numbers/"
    "service-names-port-numbers.xhtml"
)


class Clause192Report(BaseReport):
    """ITSAR Clause 1.9.2 -- Port Scanning Compliance Report (DOCX)."""

    CLAUSE_ID = "1.9.2"

    SCAN_TYPES = {
        "TC1_TCP_SCAN":  ("TCP SYN Scan",  "nmap -sS -p- -Pn -n -T4"),
        "TC2_UDP_SCAN":  ("UDP Scan",       "nmap -sU -p- -Pn -n -T4"),
        "TC3_SCTP_SCAN": ("SCTP INIT Scan", "nmap -sY -p- -Pn -n -T4"),
    }

    # ------------------------------------------------------------------ helpers
    def _add_port_classification_table(self, doc, sub_results, scan_label):
        """
        Add a table showing each discovered port, its IANA/RFC service,
        the RFC citation, and a per-port PASS/FAIL verdict.
        """
        if not sub_results:
            doc.add_paragraph("No open ports discovered for this scan type.")
            return

        self.add_itsar_subheading(
            doc, f"Port Classification — {scan_label}", 3)

        headers = ["Port", "Protocol", "nmap Service", "RFC Service",
                   "RFC / Standard", "Common", "Verdict"]
        table = doc.add_table(rows=len(sub_results) + 1, cols=len(headers))
        table.style = "Table Grid"

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self.style_table_header(cell)

        for r, sr in enumerate(sub_results, start=1):
            table.rows[r].cells[0].text = str(sr["port"])
            table.rows[r].cells[1].text = sr.get("proto", "").upper()
            table.rows[r].cells[2].text = sr.get("nmap_service", "")
            table.rows[r].cells[3].text = sr.get("rfc_service", "unknown")
            table.rows[r].cells[4].text = sr.get("rfc_url", IANA_URL)
            table.rows[r].cells[5].text = "Yes" if sr.get("is_common") else "No"

            verdict_cell = table.rows[r].cells[6]
            verdict = sr.get("status", "FAIL")
            verdict_cell.text = verdict
            for para in verdict_cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = GREEN if verdict == "PASS" else RED

        self.add_data_cell_padding(table, skip_first_row=True)
        self.prevent_table_row_split(table)

        # Citation note
        ref = doc.add_paragraph()
        ref_run = ref.add_run(
            f"Port assignments verified against the IANA Service Name and "
            f"Transport Protocol Port Number Registry: {IANA_URL}")
        ref_run.italic = True
        ref_run.font.size = Pt(8)
        ref_run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

    # ---------------------------------------------------------------- generate
    def generate(self, context, results):
        doc = Document()

        self.add_page_number(doc)
        self.add_title(doc)

        # ── Front Page ──
        self._add_front_page(doc, context, results)

        # ── 1. DUT Details ──
        self.add_dut_details(doc, context, section_num="1")
        doc.add_paragraph()

        # ── 2. ITSAR Information ──
        self.add_itsar_heading(doc, "2. ITSAR Information", 2)
        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        for i, header in enumerate(["Field", "Value"]):
            cell = table.rows[0].cells[i]
            cell.text = header
            self.style_table_header(cell)
        table.rows[1].cells[0].text = "ITSAR Section"
        table.rows[1].cells[1].text = "1.9.2"
        table.rows[2].cells[0].text = "Requirement"
        table.rows[2].cells[1].text = "Open Port Compliance"
        self.add_data_cell_padding(table)
        self.prevent_table_row_split(table)
        doc.add_paragraph()

        # ── 3. Requirement Description ──
        self.add_itsar_heading(doc, "3. Requirement Description", 2)
        doc.add_paragraph(
            "It shall be ensured that on all network interfaces, only vendor documented/"
            "identified ports on the transport layer respond to requests from outside the system. "
            "List of the identified open ports shall match the list of network services that are "
            "necessary for the operation of the CPE."
        )
        doc.add_paragraph()

        # ── 4. Preconditions ──
        self.add_itsar_heading(doc, "4. Preconditions", 2)
        for item in [
            "The tester system has network connectivity to the DUT.",
            f"The DUT IPv4 address ({getattr(context, 'dut_ip', 'N/A')}) is reachable.",
            "The tester has nmap, tcpdump, tshark, and Wireshark installed.",
            "The tester has sufficient privileges (root/sudo) for SYN, UDP, and SCTP scans.",
        ]:
            doc.add_paragraph(f"\u2022 {item}")
        doc.add_paragraph()

        # ── 5. Test Objective ──
        self.add_itsar_heading(doc, "5. Test Objective", 2)
        doc.add_paragraph(
            "To identify all open ports on the DUT using TCP SYN, UDP, and SCTP INIT "
            "scan techniques. Each discovered open port is matched against the IANA "
            "Service Name and Transport Protocol Port Number Registry (RFC-based) to "
            "determine whether it corresponds to a service commonly required for "
            "packet transfer. The test FAILS if even one port is not commonly used."
        )
        doc.add_paragraph()

        # ── 6. Test Scenario ──
        self.add_itsar_heading(doc, "6. Test Scenario", 2)

        self.add_itsar_subheading(doc, "6.1 Tools Required", 2)
        for tool in ["nmap (port scanning)", "tcpdump (packet capture)",
                     "tshark (packet analysis)", "Wireshark (visual evidence)",
                     "Linux-based tester system"]:
            doc.add_paragraph(f"\u2022 {tool}")

        self.add_itsar_subheading(doc, "6.2 Test Execution Steps", 2)
        for step in [
            "Start packet capture on the tester interface using tcpdump.",
            "Run nmap scan against the DUT for all 65535 ports.",
            "Stop the packet capture after the scan completes.",
            "Parse nmap output to identify open ports.",
            "Classify each port against the IANA/RFC well-known port registry.",
            "Take Wireshark screenshots for each discovered open port.",
            "Mark FAIL if any discovered port is not commonly used for packet transfer.",
        ]:
            doc.add_paragraph(f"\u2022 {step}")
        doc.add_paragraph()

        # ── 7. Expected Results ──
        self.add_itsar_heading(doc, "7. Expected Results", 2)
        doc.add_paragraph(
            "Only documented and necessary service ports should be found open. "
            "Every open port must map to a well-known service defined in an RFC or "
            "IANA standard that is commonly required for packet transfer on a CPE. "
            "The test FAILS if even one open port does not meet this criterion."
        )
        doc.add_paragraph()

        # ── 8. Test Execution ──
        self.add_itsar_heading(doc, "8. Test Execution", 2)

        for idx, tc in enumerate(results, start=1):
            scan_info = self.SCAN_TYPES.get(tc.name, (tc.name, "nmap"))
            scan_label, scan_cmd = scan_info

            h = self.add_itsar_subheading(doc, f"8.{idx} Test Case: {tc.name}", 2)
            self.keep_with_next(h)

            doc.add_paragraph(f"Description: {tc.description}")

            self.add_bold_paragraph(doc, "Scan Type:")
            doc.add_paragraph(scan_label)

            self.add_bold_paragraph(doc, "Execution Command:")
            doc.add_paragraph(f"{scan_cmd} {getattr(context, 'dut_ip', 'N/A')}")

            p = doc.add_paragraph("Result: ")
            run = p.add_run(tc.status)
            run.bold = True
            run.font.color.rgb = GREEN if tc.status.upper() == "PASS" else RED

            # ── Port Classification Table ──
            if hasattr(tc, "sub_results") and tc.sub_results:
                self._add_port_classification_table(
                    doc, tc.sub_results, scan_label)

                # Per-port observations
                non_std = [sr for sr in tc.sub_results if not sr.get("is_common")]
                std = [sr for sr in tc.sub_results if sr.get("is_common")]

                obs_heading = doc.add_paragraph()
                obs_run = obs_heading.add_run("Observations:")
                obs_run.bold = True
                obs_run.font.size = Pt(10)
                obs_run.font.color.rgb = PURPLE

                if non_std:
                    port_list = ", ".join(
                        f"{sr['port']}/{sr.get('proto','').upper()} "
                        f"({sr.get('rfc_service','unknown')})"
                        for sr in non_std
                    )
                    obs = doc.add_paragraph()
                    obs_text = obs.add_run(
                        f"NON-COMPLIANT: {len(non_std)} port(s) are NOT commonly "
                        f"used for packet transfer and should not be open on a CPE: "
                        f"{port_list}. Each port's defining RFC/standard is cited in "
                        f"the table above. These services represent an expanded "
                        f"attack surface and must be closed or justified by the "
                        f"vendor's operational documentation."
                    )
                    obs_text.italic = True
                    obs_text.font.size = Pt(9)
                    obs_text.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
                if std:
                    port_list = ", ".join(
                        f"{sr['port']}/{sr.get('proto','').upper()}"
                        for sr in std
                    )
                    obs2 = doc.add_paragraph()
                    obs2_text = obs2.add_run(
                        f"Compliant ports ({len(std)}): {port_list} — these are "
                        f"well-known services commonly required for CPE operation."
                    )
                    obs2_text.italic = True
                    obs2_text.font.size = Pt(9)
                    obs2_text.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

            self.add_grey_horizontal_line(doc)

            for evidence in tc.evidence:
                screenshot = evidence.get("screenshot") if isinstance(evidence, dict) else getattr(evidence, "screenshot", None)
                if screenshot and os.path.exists(screenshot):
                    self.add_screenshot_block(doc, f"Evidence: {os.path.basename(screenshot)}", screenshot)
                    doc.add_paragraph()

            self.embed_testcase_screenshots(doc, self.CLAUSE_ID, tc.name, label_prefix=f"TC{idx} -- ")

        doc.add_paragraph()

        # ── 9. Test Observation ──
        self.add_itsar_heading(doc, "9. Test Observation", 2)

        # Collect all non-standard ports across all test cases
        all_non_std = []
        all_std = []
        for tc in results:
            for sr in getattr(tc, "sub_results", []):
                if sr.get("is_common"):
                    all_std.append(sr)
                else:
                    all_non_std.append(sr)

        failed = [tc for tc in results if tc.status.upper() != "PASS"]
        if failed:
            names = ", ".join(tc.name for tc in failed)
            doc.add_paragraph(
                f"The following scan(s) reported non-compliant ports: {names}."
            )
            if all_non_std:
                port_detail = "; ".join(
                    f"Port {sr['port']}/{sr.get('proto','').upper()} — "
                    f"{sr.get('rfc_service','unknown')} "
                    f"(Ref: {sr.get('rfc_url', IANA_URL)})"
                    for sr in all_non_std
                )
                doc.add_paragraph(
                    f"Non-standard open ports discovered: {port_detail}. "
                    f"These ports are not commonly required for packet transfer "
                    f"on a CPE and their presence triggers an automatic FAIL verdict."
                )
        else:
            total_ports = len(all_std) + len(all_non_std)
            doc.add_paragraph(
                f"All port scanning test cases completed successfully. "
                f"{total_ports} open port(s) were discovered, all of which "
                f"map to well-known services commonly required for packet "
                f"transfer per IANA/RFC definitions."
            )
        doc.add_paragraph()

        # ── 10. Test Case Result Summary ──
        self.add_result_summary(doc, results, section_num="10")
        doc.add_paragraph()

        # ── 11. Compliance Analysis ──
        self._add_compliance_analysis(doc, results, section_num="11")
        doc.add_paragraph()

        # ── 12. Conclusion & Recommendations ──
        self._add_conclusion(doc, results, section_num="12", recommendations=[
            "Close all unnecessary TCP/UDP/SCTP ports on the DUT.",
            "Ensure only vendor-documented services are running.",
            "Apply firewall rules to restrict access to management interfaces.",
            "Disable any debug or development services in production firmware.",
            "Re-run this test suite after any configuration change.",
        ])

        # Save
        report_path = self.get_report_path(self.CLAUSE_ID)
        doc.save(report_path)
        return report_path
