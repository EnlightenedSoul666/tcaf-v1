import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime


# ─────────────────────────────────────────────
# COLOUR / STYLE CONSTANTS
# ─────────────────────────────────────────────
PURPLE       = RGBColor(75, 0, 130)    # #4B0082
WHITE        = RGBColor(255, 255, 255)
GREEN        = RGBColor(0, 128, 0)
RED          = RGBColor(255, 0, 0)
PURPLE_HEX   = "4B0082"
LIGHT_PURPLE = "F3ECFA"               # screenshot block background


class DOCXGenerator:

    def __init__(self, output_dir):
        self.output_dir = output_dir

    # ─────────────────────────────────────────
    # HEADING  (purple text + purple underline)
    # ─────────────────────────────────────────
    def _add_itsar_heading(self, doc, text, level=1):
        """Styled section heading with a purple bottom border."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(2)

        run = para.add_run(text)
        run.bold            = True
        run.font.size       = Pt(16 if level == 1 else 14)
        run.font.color.rgb  = PURPLE

        # bottom border underline
        pPr   = para._p.get_or_add_pPr()
        pBdr  = OxmlElement("w:pBdr")
        bot   = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "12")
        bot.set(qn("w:space"), "2")
        bot.set(qn("w:color"), PURPLE_HEX)
        pBdr.append(bot)
        pPr.append(pBdr)

        return para

    # ─────────────────────────────────────────
    # SUB-HEADING  (purple text, no underline)
    # ─────────────────────────────────────────
    def _add_itsar_subheading(self, doc, text, level=2):
        """Lighter heading for numbered sub-sections."""
        para = doc.add_heading(text, level=level)
        run = para.runs[0]
        run.bold           = True
        run.font.size      = Pt(16 if level == 1 else 14)
        run.font.color.rgb = PURPLE
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(8)
        return para

    # ─────────────────────────────────────────
    # BOLD PARAGRAPH
    # ─────────────────────────────────────────
    def _add_bold_paragraph(self, doc, text):
        p = doc.add_paragraph()
        p.add_run(text).bold = True
        return p

    # ─────────────────────────────────────────
    # TABLE HEADER CELL  (purple bg, white bold)
    # ─────────────────────────────────────────
    def _style_table_header(self, cell, color=PURPLE_HEX):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold           = True
                run.font.color.rgb = WHITE

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.top_margin    = Inches(0.15)
        cell.bottom_margin = Inches(0.15)
        cell.left_margin   = Inches(0.15)
        cell.right_margin  = Inches(0.15)

    # ─────────────────────────────────────────
    # PREVENT TABLE ROW SPLIT
    # ─────────────────────────────────────────
    def _prevent_table_row_split(self, table):
        for row in table.rows:
            trPr      = row._tr.get_or_add_trPr()
            cantSplit = OxmlElement("w:cantSplit")
            trPr.append(cantSplit)

    # ─────────────────────────────────────────
    # KEEP-WITH-NEXT
    # ─────────────────────────────────────────
    def _keep_with_next(self, para):
        para.paragraph_format.keep_with_next = True
        para.paragraph_format.keep_together  = True
        return para

    # ─────────────────────────────────────────
    # PAGE NUMBER FOOTER
    # ─────────────────────────────────────────
    def _add_page_number(self, doc):
        footer    = doc.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.clear()

        for fld_type in ("begin", None, "end"):
            r = paragraph.add_run()
            if fld_type in ("begin", "end"):
                fc = OxmlElement("w:fldChar")
                fc.set(qn("w:fldCharType"), fld_type)
                r._r.append(fc)
            else:
                it = OxmlElement("w:instrText")
                it.text = "PAGE"
                r._r.append(it)

    # ─────────────────────────────────────────
    # GREY HORIZONTAL SEPARATOR LINE
    # ─────────────────────────────────────────
    def _add_grey_horizontal_line(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)

        pPr    = p._p.get_or_add_pPr()
        pBdr   = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "BFBFBF")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ─────────────────────────────────────────
    # TITLE PAGE
    # ─────────────────────────────────────────
    def _add_title(self, doc, title_text="Telecom Compliance Automation Framework (TCAF)"):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = title.add_run(title_text)
        run.bold           = True
        run.font.size      = Pt(26)
        run.font.color.rgb = PURPLE

        doc.add_paragraph()

    # ─────────────────────────────────────────
    # FRONT PAGE  (mirrors add_front_page style)
    # ─────────────────────────────────────────
    def _add_front_page(self, doc, meta=None):
        """
        Professional front page with title, metadata table, and result banner.

        meta dict keys (all optional):
            dut_name, dut_version, os_hash, config_hash,
            start_time, end_time, final_result,
            itsar_id, itsar_version
        """
        if meta is None:
            meta = {}

        # ── Title ──────────────────────────────
        self._add_title(doc)

        # ── Subtitle ───────────────────────────
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run("Cryptographic Based Secure Communication")
        sub_run.bold       = True
        sub_run.font.size  = Pt(16)
        sub_run.font.color.rgb = PURPLE
        doc.add_paragraph()

        # ── Meta Table ─────────────────────────
        rows_data = [
            ("ITSAR ID",         meta.get("itsar_id",      "ITSAR111092401")),
            ("ITSAR Version",    meta.get("itsar_version",  "1.0")),
            ("DUT Name",         meta.get("dut_name",       "N/A")),
            ("DUT Version",      meta.get("dut_version",    "N/A")),
            ("OS Hash",          meta.get("os_hash",        "N/A")),
            ("Config Hash",      meta.get("config_hash",    "N/A")),
            ("Test Start Time",  meta.get("start_time",     datetime.now().strftime("%Y-%m-%d %H:%M:%S"))),
            ("Test End Time",    meta.get("end_time",       datetime.now().strftime("%Y-%m-%d %H:%M:%S"))),
        ]

        table = doc.add_table(rows=len(rows_data) + 1, cols=2)
        table.style = "Table Grid"

        for i, header in enumerate(["Parameter", "Value"]):
            cell      = table.rows[0].cells[i]
            cell.text = header
            self._style_table_header(cell)

        for i, (key, val) in enumerate(rows_data, start=1):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(val)

        self._add_data_cell_padding(table)
        self._prevent_table_row_split(table)

        doc.add_paragraph()

        # ── Final Result Banner ─────────────────
        final_result = meta.get("final_result", "FAIL")
        result_color = GREEN if final_result.upper() == "PASS" else RED

        banner = doc.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        banner_run = banner.add_run(f"OVERALL RESULT:  {final_result}")
        banner_run.bold           = True
        banner_run.font.size      = Pt(18)
        banner_run.font.color.rgb = result_color

        doc.add_page_break()

    # ─────────────────────────────────────────
    # SECTION 1 – Access and Authorization
    # ─────────────────────────────────────────
    def _add_access_authorization(self, doc):
        self._add_itsar_heading(doc, "1. Access and Authorization", 2)
        doc.add_paragraph(
            "This section verifies that access to the DUT is restricted to authorized "
            "entities using secure communication mechanisms."
        )

    # ─────────────────────────────────────────
    # SECTION 2 – Cryptographic Secure Communication
    # ─────────────────────────────────────────
    def _add_cryptographic_secure_comm(self, doc):
        self._add_itsar_heading(doc, "2. Cryptographic Based Secure Communication", 2)
        doc.add_paragraph(
            "The DUT shall protect data in transit using industry-standard and "
            "NIST-recommended cryptographic protocols and algorithms."
        )

    # ─────────────────────────────────────────
    # SECTION 3 – Requirement Description
    # ─────────────────────────────────────────
    def _add_requirement(self, doc):
        self._add_itsar_heading(doc, "3. Requirement Description", 2)
        doc.add_paragraph(
            "Secure communication mechanism between the Network product and the connected "
            "entities shall use only the industry standard and NIST recommended cryptographic "
            "protocols such as IPSEC, VPN, SSH, TLS/SSL, etc. Also, Network product shall "
            "provide all cryptographic service such as encryption, decryption, key exchange, "
            "authentication, data integrity etc. using the industry accepted and NIST "
            "recommended cryptographic algorithms (with standard key lengths) such as SHA, "
            "Diffie-Hellman, AES, RSA etc."
        )

    # ─────────────────────────────────────────
    # SECTION 4 – DUT Configuration
    # ─────────────────────────────────────────
    def _add_dut_configuration(self, doc, context, nmap_data=None):
        self._add_itsar_heading(doc, "4. DUT Configuration", 2)

        # 4.1 Communication Protocols
        h = self._add_itsar_subheading(doc, "4.1 Verification of Supported Communication Protocols", 2)
        self._keep_with_next(h)
        doc.add_paragraph(
            "Nmap scanning was performed to identify the communication protocols and "
            "services supported by the DUT."
        )

        if nmap_data:
            self._add_bold_paragraph(doc, "Execution Command:")
            doc.add_paragraph(nmap_data.get("user_input", "N/A"))

            self._add_bold_paragraph(doc, "Executed Command Output:")
            doc.add_paragraph(nmap_data.get("terminal_output") or "No output")

            screenshot = nmap_data.get("screenshot", "")
            if screenshot and os.path.exists(screenshot):
                self._add_screenshot_block(
                    doc,
                    "DUT Configuration : Nmap Scan Screenshot",
                    screenshot
                )

        # 4.2 DUT Details Table
        self._add_itsar_subheading(doc, "4.2 DUT Details", 2)

        rows_data = [
            ("Device",           getattr(context, "dut_model",    "N/A")),
            ("Serial Number",    getattr(context, "dut_serial",   "N/A")),
            ("Firmware Version", getattr(context, "dut_firmware", "N/A")),
            ("DUT IP Address",   getattr(context, "ssh_ip",       "N/A")),
        ]

        table = doc.add_table(rows=len(rows_data) + 1, cols=2)
        table.style = "Table Grid"

        for i, header in enumerate(["Parameter", "Value"]):
            cell      = table.rows[0].cells[i]
            cell.text = header
            self._style_table_header(cell)

        for i, (key, val) in enumerate(rows_data, start=1):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(val)

        self._add_data_cell_padding(table)
        self._prevent_table_row_split(table)

        # 4.3 SSH Host Key Configuration
        self._add_itsar_subheading(doc, "4.3 SSH Host Key Configuration", 2)
        doc.add_paragraph(
            "The client was configured to store the DUT's SSH host key in the known_hosts "
            "file, ensuring mutual authentication during SSH connection establishment."
        )

    # ─────────────────────────────────────────
    # SECTION 5 – ITSAR Information
    # ─────────────────────────────────────────
    def _add_itsar_info(self, doc, context):
        self._add_itsar_heading(doc, "5. ITSAR Information", 2)

        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"

        for i, header in enumerate(["Field", "Value"]):
            cell      = table.rows[0].cells[i]
            cell.text = header
            self._style_table_header(cell)

        table.rows[1].cells[0].text = "ITSAR Section"
        table.rows[1].cells[1].text = getattr(context, "itsar_section", "N/A")

        table.rows[2].cells[0].text = "Requirement"
        table.rows[2].cells[1].text = getattr(context, "itsar_requirement", "N/A")

        self._add_data_cell_padding(table)
        self._prevent_table_row_split(table)

    # ─────────────────────────────────────────
    # SECTION 6 – Pre-Conditions
    # ─────────────────────────────────────────
    def _add_preconditions(self, doc):
        self._add_itsar_heading(doc, "6. Preconditions", 2)
        for item in [
            "OEM should provide the list of entities that connect to DUT through which "
            "interfaces and the connection protocols between them.",
            "Network product documentation stating which security protocols for protection "
            "of data in transit are implemented and which profiles in TS 33.310 and TS 33.210 "
            "are applicable is provided by the vendor.",
        ]:
            doc.add_paragraph(f"• {item}")

    # ─────────────────────────────────────────
    # SECTION 7 – Test Objective
    # ─────────────────────────────────────────
    def _add_test_objective(self, doc):
        self._add_itsar_heading(doc, "7. Test Objective", 2)
        doc.add_paragraph(
            "To verify if the DUT management traffic shall be protected strictly using secure "
            "cryptographic controls prescribed in Table 1 of the latest document "
            "\u201cCryptographic Controls For Indian Telecom Security Assurance Requirements (ITSAR)\u201d."
        )

    # ─────────────────────────────────────────
    # SECTION 8 – Test Scenario
    # ─────────────────────────────────────────
    def _add_test_scenario(self, doc, testbed_image_path=None):
        self._add_itsar_heading(doc, "8. Test Scenario", 2)

        self._add_itsar_subheading(doc, "8.1 Number of Test Scenarios", 2)

        h = self._add_itsar_subheading(doc, "8.2 Test Bed Diagram", 2)
        self._keep_with_next(h)
        if testbed_image_path and os.path.exists(testbed_image_path):
            p = doc.add_paragraph()
            self._keep_with_next(p)
            p.add_run().add_picture(testbed_image_path, width=Inches(6.5))
        else:
            doc.add_paragraph("[Test Bed Diagram – insert image here]")

        self._add_itsar_subheading(doc, "8.3 Tools Required", 2)
        for tool in ["Wireshark", "OpenSSH", "SNMP", "Nmap", "Linux based tester system"]:
            doc.add_paragraph(f"• {tool}")

        self._add_itsar_subheading(doc, "8.4 Test Execution Steps", 2)
        for step in [
            "The tester shall establish a secure connection between the network product "
            "and the peer and verify that all protocol versions and combinations of "
            "cryptographic algorithms that are mandated by the security profile are supported "
            "by the network product, using a packet capturing tool (e.g., Wireshark).",
            "The tester should also try to attempt to connect to DUT using a deliberately "
            "unsupported cipher to check if connection to DUT is possible or not.",
        ]:
            doc.add_paragraph(f"• {step}")

    # ─────────────────────────────────────────
    # SECTION 9 – Expected Results
    # ─────────────────────────────────────────
    def _add_expected_results(self, doc):
        self._add_itsar_heading(doc, "9. Expected Results for Pass", 2)
        doc.add_paragraph(
            "Handshake of the selected management protocol is completed between DUT and "
            "Tester device and data exchanged between DUT and Tester device is encrypted "
            "using the selected encryption algorithm."
        )

    # ─────────────────────────────────────────
    # SCREENSHOT EVIDENCE BLOCK
    # ─────────────────────────────────────────
    def _add_screenshot_block(self, doc, title, image_path):
        """Lavender-background card with purple border and centred image."""
        TABLE_WIDTH = Inches(7.8)
        IMAGE_WIDTH = Inches(6.2)

        table = doc.add_table(rows=2, cols=1)
        table.alignment     = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False

        self._prevent_table_row_split(table)

        for row in table.rows:
            cell       = row.cells[0]
            cell.width = TABLE_WIDTH

            tcPr = cell._tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_PURPLE)
            tcPr.append(shd)

            cell.top_margin    = Inches(0.2)
            cell.bottom_margin = Inches(0.2)
            cell.left_margin   = Inches(0.3)
            cell.right_margin  = Inches(0.3)

        # Title cell
        title_cell = table.cell(0, 0)
        p_title    = title_cell.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._keep_with_next(p_title)

        run = p_title.add_run(title)
        run.bold           = True
        run.font.size      = Pt(11)
        run.font.color.rgb = PURPLE

        # Image cell
        img_cell = table.cell(1, 0)
        p_img    = img_cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.keep_together = True
        p_img.add_run().add_picture(image_path, width=IMAGE_WIDTH)

        # Purple border
        tblPr      = table._tbl.tblPr
        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"),   "single")
            border.set(qn("w:sz"),    "12")
            border.set(qn("w:color"), PURPLE_HEX)
            tblBorders.append(border)
        tblPr.append(tblBorders)

        return table

    # ─────────────────────────────────────────
    # NORMALIZE LIST HELPER
    # ─────────────────────────────────────────
    def _normalize_list(self, items):
        if not items:
            return ["None"]
        cleaned = [i.strip() for i in items if i.strip()]
        return cleaned if cleaned else ["None"]

    # ─────────────────────────────────────────
    # TWO-COLUMN STRONG/WEAK TABLE
    # ─────────────────────────────────────────
    def _add_strong_weak_table(self, doc, strong_items, weak_items, strong_label="Strong", weak_label="Weak"):
        h_row = [strong_label, weak_label]
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        t.cell(0, 0).text = h_row[0]
        t.cell(0, 1).text = h_row[1]

        # Style header cells
        for cell in t.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:fill"), PURPLE_HEX)
            tcPr.append(shd)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold           = True
                    run.font.color.rgb = WHITE

        row = t.add_row().cells
        row[0].text = "\n".join(self._normalize_list(strong_items))
        row[1].text = "\n".join(self._normalize_list(weak_items))
        self._prevent_table_row_split(t)
        return t

    # ─────────────────────────────────────────
    # SECTION 10 – SSH Test Execution
    # ─────────────────────────────────────────
    def _add_ssh_test_execution(self, doc, cipher_data=None, ssh_data=None, weak_cipher_result=None):
        self._add_itsar_heading(doc, "10. Test Execution For SSH", 2)

        # ── TC 10.1 ────────────────────────────
        self._add_itsar_subheading(doc, "10.1 Test Case Number: 1", 2)
        self._add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC1_DUT_SUPPORTS_SECURE_CIPHERS")

        self._add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph("DUT should support secure ciphers")

        self._add_bold_paragraph(doc, "c) Execution Steps:")
        p = doc.add_paragraph()
        p.add_run("• The tester should run the command ")
        bold_run = p.add_run("nmap --script ssh2-enum-algos <ip address>")
        bold_run.bold = True
        p.add_run(" and check for any unsupported cipher.")
        doc.add_paragraph("• Validate Results")
        doc.add_paragraph("• Capture evidence screenshots")

        if cipher_data:
            self._add_bold_paragraph(doc, "Execution Command:")
            doc.add_paragraph(cipher_data.get("user_input", "N/A"))

            self._add_bold_paragraph(doc, "Executed Command Output:")
            doc.add_paragraph(cipher_data.get("terminal_output") or "No output")

            details = cipher_data.get("details", {})

            h = self._add_itsar_subheading(doc, "10.1.1 DUT-Supported Encryption Algorithms", 2)
            self._keep_with_next(h)
            self._add_strong_weak_table(
                doc,
                details.get("encryption", {}).get("strong", []),
                details.get("encryption", {}).get("weak", []),
                "Strong Encryption", "Weak Encryption"
            )

            h = self._add_itsar_subheading(doc, "10.1.2 DUT-Supported MAC Algorithms", 2)
            self._keep_with_next(h)
            self._add_strong_weak_table(
                doc,
                details.get("mac", {}).get("strong", []),
                details.get("mac", {}).get("weak", []),
                "Strong MAC", "Weak MAC"
            )

            h = self._add_itsar_subheading(doc, "10.1.3 DUT-Supported Key Exchange Algorithms", 2)
            self._keep_with_next(h)
            self._add_strong_weak_table(
                doc,
                details.get("kex", {}).get("strong", []),
                details.get("kex", {}).get("weak", []),
                "Strong KEX", "Weak KEX"
            )

            h = self._add_itsar_subheading(doc, "10.1.4 DUT-Supported Host Key Algorithms", 2)
            self._keep_with_next(h)
            self._add_strong_weak_table(
                doc,
                details.get("host_key", {}).get("strong", []),
                details.get("host_key", {}).get("weak", []),
                "Strong Host Key", "Weak Host Key"
            )

            doc.add_paragraph()

            screenshot = cipher_data.get("screenshot", "")
            if screenshot and os.path.exists(screenshot):
                self._add_screenshot_block(
                    doc,
                    "Test Case 1 : Executed cmd output: sudo nmap --script ssh2-enum-algos",
                    screenshot
                )

        # ── TC 10.2 ────────────────────────────
        self._add_itsar_subheading(doc, "10.2 Test Case Number: 2", 2)
        self._add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC1_PROTECT_DATA_INFO_TRANSFER_USING_SSH")

        self._add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "DUT and tester device should support traffic protection through SSH "
            "only through provided cryptographic methods"
        )

        self._add_bold_paragraph(doc, "c) Execution Steps:")
        doc.add_paragraph(
            "• The tester should attempt to ssh into the DUT and capture the SSH "
            "traffic using wireshark"
        )
        doc.add_paragraph("• Analyze the ssh handshake in the p-cap file.")

        if ssh_data:
            self._add_bold_paragraph(doc, "Execution Command:")
            doc.add_paragraph(ssh_data.get("user_input", "N/A"))

            h = self._add_itsar_subheading(
                doc, "10.2.1 Security Ciphers Utilized During Data Transmission (Verbose Fetch)", 2
            )
            self._keep_with_next(h)

            crypto = ssh_data.get("crypto_details", {})
            enc = doc.add_table(rows=4, cols=2)
            enc.style = "Table Grid"
            enc.cell(0, 0).text = "Protocol"
            enc.cell(0, 1).text = crypto.get("protocol", "Not Found")
            enc.cell(1, 0).text = "Encryption Algorithm"
            enc.cell(1, 1).text = crypto.get("cipher", "Not Found")
            enc.cell(2, 0).text = "Key Exchange Algorithm"
            enc.cell(2, 1).text = crypto.get("kex", "Not Found")
            enc.cell(3, 0).text = "Host Key Algorithm"
            enc.cell(3, 1).text = crypto.get("host_key", "Not Found")

            doc.add_paragraph()
            self._keep_with_next(doc.add_paragraph())

            screenshots = ssh_data.get("screenshots", [])
            if screenshots:
                kex_algo    = crypto.get("kex", "Unknown")
                cipher_algo = crypto.get("cipher", "Unknown")
                nist        = ssh_data.get("nist_validation", {})
                kex_label   = "Secure" if nist.get("kex") == "PASS" else "Insecure"
                enc_label   = "Secure" if nist.get("encryption") == "PASS" else "Insecure"

                titles = [
                    "Test Case 2 : SSH CLI Packet Capture showing SSH Handshake Traffic",
                    f"Test Case 2 : SSH Handshake showing {kex_label} Key Exchange Algorithm : {kex_algo}",
                    f"Test Case 2 : SSH Data Encryption using {enc_label} Cipher : {cipher_algo}",
                ]
                overviews = [
                    "The above screenshot shows the command line packet capture using t-shark while "
                    "performing SSH communication between the tester system and the DUT. The captured "
                    "packets include the SSH handshake packets exchanged during the establishment of the SSH session.",
                    f"The above screenshot shows the SSH handshake process where the Key Exchange algorithm "
                    f"{kex_algo} is used to securely establish the cryptographic session. According to "
                    f"ITSAR guidelines, this key exchange algorithm is classified as {kex_label}.",
                    f"The above screenshot shows the encrypted SSH communication packets after the successful "
                    f"handshake. The Encryption algorithm {cipher_algo} is used to encrypt the SSH data packets. "
                    f"According to ITSAR guidelines, this cipher is classified as {enc_label}.",
                ]

                for idx, img in enumerate(screenshots):
                    if os.path.exists(img):
                        self._add_screenshot_block(doc, titles[idx], img)
                        self._add_itsar_heading(doc, f"10.2.{2 + idx}. Overview", 2)
                        doc.add_paragraph(overviews[idx])
                        spacer = doc.add_paragraph()
                        spacer.paragraph_format.space_after = Pt(18)

        # ── Weak cipher negative tests ──────────
        if weak_cipher_result and weak_cipher_result.get("screenshots"):
            results     = weak_cipher_result.get("results", [])
            screenshots = weak_cipher_result.get("screenshots", [])

            for idx, (res, img) in enumerate(zip(results, screenshots)):
                algo       = res.get("algorithm", "Unknown")
                algo_type  = res.get("type", "Unknown")
                negotiated = res.get("negotiated", False)
                neg_text   = "successfully negotiated" if negotiated else "rejected by the DUT"

                type_labels = {
                    "cipher":   (f"SSH Weak Cipher Attempt : {algo}",       "encryption algorithm"),
                    "mac":      (f"SSH Weak MAC Attempt : {algo}",           "MAC algorithm"),
                    "kex":      (f"SSH Weak Key Exchange Attempt : {algo}",  "key exchange algorithm"),
                    "host_key": (f"SSH Weak Host Key Attempt : {algo}",      "host key algorithm"),
                }
                title, algo_label = type_labels.get(algo_type, (f"SSH Weak Algorithm Attempt : {algo}", "algorithm"))

                if os.path.exists(img):
                    self._add_screenshot_block(doc, title, img)
                    self._add_itsar_heading(doc, f"10.2.{5 + idx}. Overview", 2)
                    doc.add_paragraph(
                        f"The above screenshot shows the SSH negotiation attempt where the weak "
                        f"{algo_label} {algo} was forced during SSH communication. The algorithm was "
                        f"{neg_text}. According to ITSAR guidelines, this algorithm is classified as Insecure."
                    )
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_after = Pt(18)

    # ─────────────────────────────────────────
    # SECTION 11 – SSH Test Observation
    # ─────────────────────────────────────────
    def _add_ssh_observation(self, doc, cipher_data, ssh_data, weak_cipher_result):
        tc1 = cipher_data.get("result", "FAIL") if cipher_data else "FAIL"
        tc2 = ssh_data.get("final_result", "FAIL") if ssh_data else "FAIL"

        tc3 = "PASS"
        weak_results = weak_cipher_result.get("results", []) if weak_cipher_result else []
        for r in weak_results:
            if r.get("negotiated") is True:
                tc3 = "FAIL"
                break

        failed = []
        if tc1 == "FAIL": failed.append("secure SSH cipher support (TC1)")
        if tc2 == "FAIL": failed.append("secure SSH communication protection (TC2-PART-1)")
        if tc3 == "FAIL": failed.append("weak SSH algorithm negotiation protection (TC2-PART-2)")

        if failed:
            observation = (
                "It was observed that the Device Under Test (DUT) does not fully comply with "
                "the prescribed secure cryptographic requirements, as one or more security "
                "validation test cases have failed. The failure was identified in: "
                + ", ".join(failed)
                + ". This indicates that the DUT permits insecure cryptographic configurations "
                "which may weaken the security of SSH communication."
            )
        else:
            observation = (
                "It was observed that the Device Under Test (DUT) complies with the prescribed "
                "secure cryptographic requirements. All test cases related to secure SSH cipher "
                "support, encrypted communication protection, and weak algorithm negotiation "
                "have successfully passed."
            )

        self._add_itsar_heading(doc, "11. Test Observation for SSH", 2)
        doc.add_paragraph(observation)

        return tc1, tc2, tc3

    # ─────────────────────────────────────────
    # SECTION 12 – SSH Result Table
    # ─────────────────────────────────────────
    def _add_ssh_result_table(self, doc, tc1, tc2, tc3, weak_results=None):
        if weak_results is None:
            weak_results = []

        h = self._add_itsar_heading(doc, "12. Test Case Result for SSH", 2)
        self._keep_with_next(h)

        headers = ["SL. No", "TEST CASE NAME", "PASS/FAIL", "Remarks"]
        rt = doc.add_table(rows=4, cols=4)
        rt.style = "Table Grid"

        for i, header in enumerate(headers):
            cell = rt.cell(0, i)
            cell.text = ""
            p    = cell.paragraphs[0]
            run  = p.add_run(header)
            run.bold           = True
            run.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._style_table_header(cell, PURPLE_HEX)

        rt.cell(1, 0).text = "1"
        rt.cell(1, 1).text = "TC1_DUT_SUPPORTS_SECURE_CIPHERS"
        rt.cell(1, 2).text = tc1
        rt.cell(1, 3).text = (
            "All cryptographic algorithms comply with policy." if tc1 == "PASS"
            else "One or more weak cryptographic algorithms detected."
        )

        rt.cell(2, 0).text = "2"
        rt.cell(2, 1).text = "TC2_PROTECT_DATA_INFO_TRANSFER_USING_SSH_PART_1"
        rt.cell(2, 2).text = tc2
        rt.cell(2, 3).text = (
            "SSH communication is protected using secure cryptographic methods." if tc2 == "PASS"
            else "SSH communication used weak or unsupported cryptographic methods."
        )

        weak_algos = [r["algorithm"] for r in weak_results if r.get("negotiated")]
        rt.cell(3, 0).text = "3"
        rt.cell(3, 1).text = "TC2_SSH_WEAK_ALGORITHM_NEGOTIATION_PART_2"
        rt.cell(3, 2).text = tc3
        rt.cell(3, 3).text = (
            "The DUT correctly rejects weak SSH algorithms during negotiation." if tc3 == "PASS"
            else f"The DUT allows negotiation of weak SSH algorithms: {', '.join(weak_algos)}."
        )

        # Colour PASS/FAIL cells
        for row in rt.rows[1:]:
            status_cell = row.cells[2]
            status_text = status_cell.text.strip()
            color = GREEN if status_text == "PASS" else RED
            for para in status_cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = color

        for row in rt.rows[1:]:
            for cell in row.cells:
                cell.top_margin    = Inches(0.12)
                cell.bottom_margin = Inches(0.12)
                cell.left_margin   = Inches(0.12)
                cell.right_margin  = Inches(0.12)

        self._prevent_table_row_split(rt)

    # ─────────────────────────────────────────
    # SECTION 13 – HTTPS Test Execution
    # ─────────────────────────────────────────
    def _add_https_test_execution(self, doc, https_cipher_data=None, https_data=None):
        self._add_itsar_heading(doc, "13. Test Execution For HTTPS", 2)

        # ── TC 13.1 ────────────────────────────
        self._add_itsar_subheading(doc, "13.1 Test Case Number: 1", 2)
        self._add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC1_HTTPS_CRYPTO_HARDENING")

        self._add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph("DUT should support only strong TLS cryptographic ciphers for HTTPS communication")

        self._add_bold_paragraph(doc, "c) Execution Steps:")
        p = doc.add_paragraph()
        p.add_run("• The tester should run the command ")
        bold_run = p.add_run("nmap --script ssl-enum-ciphers -p 443 <ip address>")
        bold_run.bold = True
        p.add_run(" to enumerate the TLS ciphers supported by the DUT.")
        doc.add_paragraph("• Validate that only strong TLS ciphers are supported.")
        doc.add_paragraph("• Capture evidence screenshots.")

        if https_cipher_data:
            self._add_bold_paragraph(doc, "Execution Command:")
            doc.add_paragraph(https_cipher_data.get("user_input", "N/A"))

            self._add_bold_paragraph(doc, "Executed Command Output:")
            doc.add_paragraph(https_cipher_data.get("terminal_output") or "No output")

            details = https_cipher_data.get("details", {})

            for version_label, section_prefix in [("TLSv1.2", "13.1"), ("TLSv1.3", "13.1")]:
                ver = details.get(version_label, {})
                sub_offset = 1 if version_label == "TLSv1.2" else 4

                h = self._add_itsar_subheading(
                    doc, f"{section_prefix}.{sub_offset} DUT-Supported {version_label} Encryption Algorithms", 2
                )
                self._keep_with_next(h)
                self._add_strong_weak_table(
                    doc,
                    ver.get("encryption", {}).get("strong", []),
                    ver.get("encryption", {}).get("weak", []),
                    "Strong Encryption", "Weak Encryption"
                )

                h = self._add_itsar_subheading(
                    doc, f"{section_prefix}.{sub_offset + 1} DUT-Supported {version_label} MAC Algorithms", 2
                )
                self._keep_with_next(h)
                self._add_strong_weak_table(
                    doc,
                    ver.get("mac", {}).get("strong", []),
                    ver.get("mac", {}).get("weak", []),
                    "Strong MAC", "Weak MAC"
                )

                h = self._add_itsar_subheading(
                    doc, f"{section_prefix}.{sub_offset + 2} DUT-Supported {version_label} Key Exchange Algorithms", 2
                )
                self._keep_with_next(h)
                self._add_strong_weak_table(
                    doc,
                    ver.get("kex", {}).get("strong", []),
                    ver.get("kex", {}).get("weak", []),
                    "Strong KEX", "Weak KEX"
                )

            doc.add_paragraph()

            screenshot = https_cipher_data.get("screenshot", "")
            if screenshot and os.path.exists(screenshot):
                self._add_screenshot_block(
                    doc,
                    "Test Case 1 : Executed cmd output: nmap --script ssl-enum-ciphers -p 443",
                    screenshot
                )

        # ── TC 13.2 ────────────────────────────
        self._add_itsar_subheading(doc, "13.2 Test Case Number: 2", 2)
        self._add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC2_PROTECT_DATA_INFO_TRANSFER_USING_HTTPS")

        self._add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "DUT and tester devices should support traffic protection through HTTPS "
            "using secure TLS cryptographic protocols."
        )

        self._add_bold_paragraph(doc, "c) Execution Steps:")
        doc.add_paragraph("• The tester should initiate HTTPS communication with the DUT using openssl.")
        doc.add_paragraph("• Capture the TLS handshake packets using t-shark.")
        doc.add_paragraph(
            "• Analyze the Server Hello packet in the captured pcap file to identify "
            "the TLS protocol version and cipher used during communication."
        )

        if https_data:
            self._add_bold_paragraph(doc, "Execution Command:")
            doc.add_paragraph(https_data.get("user_input", "N/A"))

            self._add_bold_paragraph(doc, "Executed Command Output:")
            doc.add_paragraph(https_data.get("terminal_output") or "No output")

            h = self._add_itsar_subheading(
                doc, "13.2.1 Security Ciphers Utilized During HTTPS Data Transmission", 2
            )
            self._keep_with_next(h)

            crypto = https_data.get("crypto_details", {})
            tls_table = doc.add_table(rows=2, cols=2)
            tls_table.style = "Table Grid"
            tls_table.cell(0, 0).text = "Protocol"
            tls_table.cell(0, 1).text = crypto.get("protocol", "Not Found")
            tls_table.cell(1, 0).text = "Encryption Algorithm"
            tls_table.cell(1, 1).text = crypto.get("cipher", "Not Found")
            self._prevent_table_row_split(tls_table)

            doc.add_paragraph()

            screenshots = https_data.get("screenshots", [])
            if screenshots:
                protocol     = crypto.get("protocol", "Unknown")
                cipher       = crypto.get("cipher", "Unknown")
                nist         = https_data.get("nist_validation", {})
                cipher_label = "Secure" if nist.get("cipher") == "PASS" else "Insecure"

                titles = [
                    "Test Case 2 : HTTPS CLI Packet Capture showing TLS Handshake Traffic",
                    f"Test Case 2 : TLS Server Hello showing {cipher_label} Cipher : {cipher}",
                ]
                overviews = [
                    "The above screenshot shows the command line packet capture using t-shark while "
                    "performing HTTPS communication. The captured packets include the TLS handshake "
                    "packets exchanged during the secure HTTPS session establishment.",
                    f"The above screenshot shows the TLS Server Hello packet captured during the HTTPS "
                    f"handshake. The Server Hello confirms that TLS protocol version {protocol} and "
                    f"cipher suite {cipher} are used. According to ITSAR guidelines, the cipher "
                    f"{cipher} is classified as {cipher_label}.",
                ]

                for idx, img in enumerate(screenshots):
                    if os.path.exists(img):
                        self._add_screenshot_block(doc, titles[idx], img)
                        self._add_itsar_heading(doc, f"13.2.{2 + idx}. Overview", 2)
                        doc.add_paragraph(overviews[idx])
                        spacer = doc.add_paragraph()
                        spacer.paragraph_format.space_after = Pt(18)

    # ─────────────────────────────────────────
    # SECTION 14 – HTTPS Observation
    # ─────────────────────────────────────────
    def _add_https_observation(self, doc, https_cipher_data, https_data):
        tc1 = https_cipher_data.get("result", "FAIL") if https_cipher_data else "FAIL"
        tc2 = https_data.get("final_result", "FAIL") if https_data else "FAIL"

        failed = []
        if tc1 == "FAIL": failed.append("secure HTTPS TLS cipher support (TC1)")
        if tc2 == "FAIL": failed.append("secure HTTPS communication protection (TC2)")

        if failed:
            observation = (
                "It was observed that the Device Under Test (DUT) does not fully comply with "
                "the prescribed secure cryptographic requirements for HTTPS communication. "
                "The failure was identified in: "
                + ", ".join(failed)
                + ". This indicates that the DUT permits insecure TLS protocol versions or "
                "cipher configurations which may weaken the security of HTTPS communication."
            )
        else:
            observation = (
                "It was observed that the Device Under Test (DUT) complies with the prescribed "
                "secure cryptographic requirements for HTTPS communication. All test cases related "
                "to secure TLS cipher support and encrypted HTTPS communication protection have "
                "successfully passed."
            )

        self._add_itsar_heading(doc, "14. Test Observation for HTTPS", 2)
        doc.add_paragraph(observation)

        return tc1, tc2

    # ─────────────────────────────────────────
    # SECTION 15 – HTTPS Result Table
    # ─────────────────────────────────────────
    def _add_https_result_table(self, doc, tc1, tc2):
        h = self._add_itsar_heading(doc, "15. Test Case Result for HTTPS", 2)
        self._keep_with_next(h)

        headers = ["SL. No", "TEST CASE NAME", "PASS/FAIL", "Remarks"]
        rt = doc.add_table(rows=3, cols=4)
        rt.style = "Table Grid"

        for i, header in enumerate(headers):
            cell = rt.cell(0, i)
            cell.text = ""
            p    = cell.paragraphs[0]
            run  = p.add_run(header)
            run.bold           = True
            run.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._style_table_header(cell, PURPLE_HEX)

        rt.cell(1, 0).text = "1"
        rt.cell(1, 1).text = "TC1_HTTPS_CRYPTO_HARDENING"
        rt.cell(1, 2).text = tc1
        rt.cell(1, 3).text = (
            "HTTPS service supports only secure TLS cryptographic algorithms." if tc1 == "PASS"
            else "Weak or unsupported TLS cryptographic algorithms detected."
        )

        rt.cell(2, 0).text = "2"
        rt.cell(2, 1).text = "TC2_PROTECT_DATA_INFO_TRANSFER_USING_HTTPS"
        rt.cell(2, 2).text = tc2
        rt.cell(2, 3).text = (
            "HTTPS communication is protected using secure TLS protocol and cipher." if tc2 == "PASS"
            else "HTTPS communication used weak or unsupported TLS protocol or cipher."
        )

        # Colour PASS/FAIL cells
        for row in rt.rows[1:]:
            status_cell = row.cells[2]
            status_text = status_cell.text.strip()
            color = GREEN if status_text == "PASS" else RED
            for para in status_cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = color

        for row in rt.rows[1:]:
            for cell in row.cells:
                cell.top_margin    = Inches(0.12)
                cell.bottom_margin = Inches(0.12)
                cell.left_margin   = Inches(0.12)
                cell.right_margin  = Inches(0.12)

        self._prevent_table_row_split(rt)

    # ─────────────────────────────────────────
    # INTERNAL HELPER – cell padding for data rows
    # ─────────────────────────────────────────
    def _add_data_cell_padding(self, table, skip_first_row=True):
        start = 1 if skip_first_row else 0
        for row in table.rows[start:]:
            for cell in row.cells:
                cell.top_margin    = Inches(0.12)
                cell.bottom_margin = Inches(0.12)
                cell.left_margin   = Inches(0.12)
                cell.right_margin  = Inches(0.12)

    # ─────────────────────────────────────────
    # GENERATE REPORT  (main entry point)
    # ─────────────────────────────────────────
    def generate(
        self,
        context,
        results,
        meta=None,
        nmap_data=None,
        cipher_data=None,
        ssh_data=None,
        weak_cipher_result=None,
        https_cipher_data=None,
        https_data=None,
        testbed_image_path=None,
    ):
        """
        Generate the full ITSAR compliance report.

        Required:
            context  – object with dut_model, dut_serial, dut_firmware, ssh_ip,
                       itsar_section, itsar_requirement
            results  – list of objects with .name, .description, .status, .evidence,
                       optionally .remarks

        Optional keyword data dicts (same structure as the second script):
            meta, nmap_data, cipher_data, ssh_data, weak_cipher_result,
            https_cipher_data, https_data, testbed_image_path
        """
        report_path = os.path.join(self.output_dir, "tcaf_report.docx")
        doc = Document()

        self._add_page_number(doc)

        # ── Front Page ──────────────────────────
        if meta:
            self._add_front_page(doc, meta)
        else:
            self._add_title(doc)

        # ── Sections 1–9 ────────────────────────
        self._add_access_authorization(doc)
        doc.add_paragraph()

        self._add_cryptographic_secure_comm(doc)
        doc.add_paragraph()

        self._add_requirement(doc)
        doc.add_paragraph()

        self._add_dut_configuration(doc, context, nmap_data)
        doc.add_paragraph()

        self._add_itsar_info(doc, context)
        doc.add_paragraph()

        self._add_preconditions(doc)
        doc.add_paragraph()

        self._add_test_objective(doc)
        doc.add_paragraph()

        self._add_test_scenario(doc, testbed_image_path)
        doc.add_paragraph()

        self._add_expected_results(doc)
        doc.add_paragraph()

        # ── SSH Sections 10–12 ──────────────────
        self._add_ssh_test_execution(doc, cipher_data, ssh_data, weak_cipher_result)
        doc.add_paragraph()

        ssh_tc1, ssh_tc2, ssh_tc3 = self._add_ssh_observation(
            doc, cipher_data, ssh_data, weak_cipher_result
        )
        doc.add_paragraph()

        weak_results = weak_cipher_result.get("results", []) if weak_cipher_result else []
        self._add_ssh_result_table(doc, ssh_tc1, ssh_tc2, ssh_tc3, weak_results)
        doc.add_paragraph()

        # ── HTTPS Sections 13–15 ────────────────
        self._add_https_test_execution(doc, https_cipher_data, https_data)
        doc.add_paragraph()

        https_tc1, https_tc2 = self._add_https_observation(doc, https_cipher_data, https_data)
        doc.add_paragraph()

        self._add_https_result_table(doc, https_tc1, https_tc2)

        doc.save(report_path)
        return report_path